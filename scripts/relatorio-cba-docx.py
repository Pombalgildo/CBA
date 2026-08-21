#!/usr/bin/env python3
"""
Relatório de Prestação de Contas e Entrega Técnica — CBA
Gera um documento Word (.docx) profissional com as cores da CBA.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ─── CBA Color Palette ───
CBA_GREEN = RGBColor(0x1a, 0x3a, 0x2a)
CBA_RED   = RGBColor(0xb9, 0x1c, 0x1c)
CBA_GOLD  = RGBColor(0xc4, 0x96, 0x2e)
CBA_DARK  = RGBColor(0x23, 0x22, 0x20)
CBA_MUTED = RGBColor(0x6b, 0x6b, 0x6b)
WHITE     = RGBColor(0xff, 0xff, 0xff)
LIGHT_BG  = "F5F3F0"
SALDO_BG  = "FEF3C7"
GREEN_BG  = "DCFCE7"
BORDER    = "D4D0C8"

# ─── Helpers ───
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="D4D0C8", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_para_border(paragraph, position='bottom', color='C4962E', sz='12'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    b = OxmlElement(f'w:{position}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    p_bdr.append(b)
    p_pr.append(p_bdr)

def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Set East Asia font too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)

def add_paragraph_with_runs(doc_or_cell, runs_data, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_after=6, line_spacing=1.3, first_line_indent=None):
    """runs_data: list of dicts {text, bold, color, size, name}"""
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    for rd in runs_data:
        r = p.add_run(rd.get('text', ''))
        set_run_font(r,
                     name=rd.get('name', 'Calibri'),
                     size=rd.get('size', 11),
                     bold=rd.get('bold', False),
                     color=rd.get('color', CBA_DARK))
    return p

def add_heading(doc, text, level=1):
    """Add a heading with CBA green color and gold underline divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
        size = 15
    else:
        pf.space_before = Pt(12)
        pf.space_after = Pt(4)
        size = 12
    pf.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, name='Cambria', size=size, bold=True, color=CBA_GREEN)
    # Set as heading style for navigation/TOC
    p.style = doc.styles[f'Heading {level}']
    # Re-apply our font (heading style overrides)
    for run in p.runs:
        set_run_font(run, name='Cambria', size=size, bold=True, color=CBA_GREEN)
    if level == 1:
        # Add gold divider
        div = doc.add_paragraph()
        div.paragraph_format.space_before = Pt(0)
        div.paragraph_format.space_after = Pt(8)
        add_para_border(div, 'bottom', 'C4962E', '12')
    return p

def add_table_styled(doc, headers, rows, col_widths_pct, saldo_row_idx=None, green_row_idx=None):
    """Create a styled table with CBA branding."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set column widths (in Cm) — A4 content width ~16cm
    total_width_cm = 16.0
    for i, w in enumerate(col_widths_pct):
        for cell in table.columns[i].cells:
            cell.width = Cm(total_width_cm * w / 100.0)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        set_run_font(r, name='Calibri', size=10, bold=True, color=WHITE)
        set_cell_bg(cell, '1A3A2A')
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        is_saldo = (saldo_row_idx is not None and ri == saldo_row_idx)
        is_green = (green_row_idx is not None and ri == green_row_idx)
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 or len(headers) == 2 and ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            # Right-align numeric columns (last columns when >2 cols)
            if len(headers) >= 3 and ci >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif len(headers) == 4 and ci == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(val)
            bold = is_saldo
            set_run_font(r, name='Calibri', size=10, bold=bold,
                         color=CBA_DARK if not is_saldo else CBA_DARK)
            # Background
            if is_saldo:
                set_cell_bg(cell, SALDO_BG)
            elif is_green:
                set_cell_bg(cell, GREEN_BG)
            elif ri % 2 == 1:
                set_cell_bg(cell, LIGHT_BG)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add some padding via cell margins
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement('w:tcMar')
            for side in ('top', 'left', 'bottom', 'right'):
                m = OxmlElement(f'w:{side}')
                m.set(qn('w:w'), '120' if side in ('left', 'right') else '60')
                m.set(qn('w:type'), 'dxa')
                tc_mar.append(m)
            tc_pr.append(tc_mar)
    return table

def add_bullet(doc, text_runs, indent_cm=0.6):
    """text_runs: list of dicts; adds a bullet paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(indent_cm)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.3
    r = p.add_run('•  ')
    set_run_font(r, name='Calibri', size=10.5, color=CBA_GOLD, bold=True)
    for rd in text_runs:
        r = p.add_run(rd.get('text', ''))
        set_run_font(r,
                     name=rd.get('name', 'Calibri'),
                     size=rd.get('size', 10.5),
                     bold=rd.get('bold', False),
                     color=rd.get('color', CBA_DARK))
    return p

# ─── Build Document ───
output_path = '/home/z/my-project/download/Relatorio_Prestacao_Contas_CBA.docx'

doc = Document()

# Page setup — A4
for section in doc.sections:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Set default document font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════

# Top gold line
top_line = doc.add_paragraph()
top_line.paragraph_format.space_before = Pt(0)
top_line.paragraph_format.space_after = Pt(60)
add_para_border(top_line, 'bottom', 'C4962E', '24')

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('Relatório de Prestação de Contas')
set_run_font(r, name='Cambria', size=22, bold=True, color=CBA_GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run('e Entrega Técnica')
set_run_font(r, name='Cambria', size=22, bold=True, color=CBA_GREEN)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Construção e Lançamento do Site Oficial')
set_run_font(r, name='Calibri', size=13, color=CBA_DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
r = p.add_run('da Convenção Baptista de Angola')
set_run_font(r, name='Calibri', size=13, color=CBA_DARK)

# Middle gold line
mid_line = doc.add_paragraph()
mid_line.paragraph_format.space_after = Pt(60)
add_para_border(mid_line, 'bottom', 'C4962E', '12')

# Recipient
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Apresentado à Direcção da')
set_run_font(r, name='Calibri', size=11, color=CBA_MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(120)
r = p.add_run('Convenção Baptista de Angola')
set_run_font(r, name='Cambria', size=12, bold=True, color=CBA_GREEN)

# Author and date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Elaborado por: Hermenegildo José Pombal')
set_run_font(r, name='Calibri', size=10, color=CBA_DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
r = p.add_run('Luanda, Julho de 2026')
set_run_font(r, name='Calibri', size=10, color=CBA_MUTED)

# Bottom red bar (using paragraph border)
bot_line = doc.add_paragraph()
bot_line.paragraph_format.space_before = Pt(60)
bot_line.paragraph_format.space_after = Pt(0)
add_para_border(bot_line, 'bottom', 'B91C1C', '36')

# Page break to body
p = doc.add_paragraph()
p.add_run().add_break(WD_BREAK.PAGE)

# ════════════════════════════════════════════
# SECTION 1: INTRODUÇÃO
# ════════════════════════════════════════════
add_heading(doc, '1. Introdução e Objectivo', level=1)

add_paragraph_with_runs(doc, [
    {'text': 'O presente relatório tem por objectivo apresentar à Direcção da Convenção Baptista de Angola (CBA) '
             'a prestação de contas referente aos fundos depositados para a aquisição do domínio do site oficial '
             'da instituição, bem como a entrega formal e conclusão do projecto de construção da plataforma digital.'}
])

add_paragraph_with_runs(doc, [
    {'text': 'Foi depositado na conta do desenvolvedor, Hermenegildo José Pombal, o montante de '},
    {'text': '100.000,00 Kz', 'bold': True},
    {'text': ' (cem mil Kwanzas) com a finalidade de adquirir o domínio de internet para o site da CBA. '
             'Após análise comparativa entre as opções de domínios disponíveis no mercado nacional angolano (.ao) '
             'e internacional (.org), optou-se pela aquisição do domínio '},
    {'text': 'cbaangola.org', 'bold': True},
    {'text': ' junto da empresa Hostinger, por razões de acessibilidade de custos e facilidade de pagamento.'}
])

add_paragraph_with_runs(doc, [
    {'text': 'Este documento detalha a utilização financeira dos recursos, a justificação técnica da escolha do '
             'domínio, a descrição completa do trabalho entregue e a proposta de devolução do saldo remanescente.'}
])

# ════════════════════════════════════════════
# SECTION 2: PRESTAÇÃO DE CONTAS
# ════════════════════════════════════════════
add_heading(doc, '2. Prestação de Contas Financeira', level=1)

add_paragraph_with_runs(doc, [
    {'text': 'Abaixo apresenta-se o detalhe financeiro dos valores recebidos, convertidos e utilizados para a '
             'aquisição do domínio. O montante original em Kwanzas foi convertido em Euros para possibilitar a '
             'compra junto de um fornecedor internacional, dada a impossibilidade de pagamento em Kwanzas junto '
             'da maioria dos fornecedores internacionais.'}
])

# Financial table
doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_table_styled(
    doc,
    headers=['Descrição', 'Valor (Kz)', 'Valor (EUR)'],
    rows=[
        ['Valor recebido (depósito da CBA)', '100.000,00', '—'],
        ['Conversão para EUR (taxa aproximada)', '100.000,00', '≈ 101,00 €'],
        ['Compra do domínio cbaangola.org (Hostinger)', '85.639,00', '≈ 86,49 €'],
        ['Saldo remanescente', '14.361,00', '≈ 14,51 €'],
    ],
    col_widths_pct=[50, 25, 25],
    saldo_row_idx=3,
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_paragraph_with_runs(doc, [
    {'text': 'Nota: ', 'bold': True},
    {'text': 'Os valores em EUR são aproximados, sujeitos à taxa de câmbio do dia da transação. '
             'O saldo remanescente de '},
    {'text': '14.361,00 Kz', 'bold': True},
    {'text': ' está disponível para devolução à CBA, conforme detalhado na Secção 5 deste relatório.'}
])

# ════════════════════════════════════════════
# SECTION 3: ANÁLISE .ao vs .org
# ════════════════════════════════════════════
add_heading(doc, '3. Análise Comparativa: Domínios .ao vs .org', level=1)

add_paragraph_with_runs(doc, [
    {'text': 'Antes da aquisição do domínio, foi realizada uma análise comparativa entre as opções de registos '
             'de domínios disponíveis no mercado nacional angolano (.ao) e no mercado internacional (.org). '
             'Abaixo apresentam-se os fornecedores angolanos consultados, com respectivos preços e formas de '
             'pagamento, bem como a comparação com a opção internacional escolhida.'}
])

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_table_styled(
    doc,
    headers=['Fornecedor', 'Domínio', 'Preço/ano', 'Pagamento'],
    rows=[
        ['DNS Angola (dnsao.com)', '.ao', '~15.000 Kz', 'Multicaixa, Transferência'],
        ['Webtica (webtica.ao)', '.ao', '~18.000 Kz', 'Multicaixa Express, BA Directo'],
        ['NBits (nbits.ao)', '.ao', '~20.000 Kz', 'Multicaixa, Transferência'],
        ['Angola Online (angolaonline.net)', '.ao', '~25.000 Kz', 'Multicaixa, Numerário'],
        ['Hostinger Angola (hostinger.ao)', '.ao', '~22.000 Kz', 'Multicaixa Express'],
        ['Hostinger (hostinger.com)', '.org', '~8.980 Kz', 'Cartão Visa/EUR'],
    ],
    col_widths_pct=[28, 12, 25, 35],
    green_row_idx=5,
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_paragraph_with_runs(doc, [
    {'text': 'Justificação da escolha: ', 'bold': True},
    {'text': 'O domínio '},
    {'text': 'cbaangola.org', 'bold': True},
    {'text': ' foi seleccionado pelos seguintes motivos:'}
])

for point in [
    'Custo significativamente inferior (~8.980 Kz/ano vs ~15.000-25.000 Kz/ano para .ao);',
    'O domínio .org é internacionalmente reconhecido como adequado para organizações e instituições;',
    'A extensão .ao, embora nacional, apresentava custos superiores e, em alguns fornecedores, processos de registo mais demorados;',
    'O alojamento web (gratuito na Vercel) é compatível com qualquer extensão de domínio.',
]:
    add_bullet(doc, [{'text': point, 'size': 10.5}])

# ════════════════════════════════════════════
# SECTION 4: ENTREGA TÉCNICA
# ════════════════════════════════════════════
p = doc.add_paragraph()
p.add_run().add_break(WD_BREAK.PAGE)

add_heading(doc, '4. Entrega Técnica do Projecto', level=1)

add_paragraph_with_runs(doc, [
    {'text': 'O site oficial da Convenção Baptista de Angola foi construído utilizando tecnologias modernas e '
             'padrões internacionais de desenvolvimento web. O projecto encontra-se '},
    {'text': 'concluído, testado e prontamente funcional', 'bold': True},
    {'text': ', acessível através do endereço '},
    {'text': 'https://www.cbaangola.org', 'bold': True},
    {'text': '.'}
])

add_heading(doc, '4.1 Tecnologias Utilizadas', level=2)

add_table_styled(
    doc,
    headers=['Componente', 'Tecnologia', 'Custo'],
    rows=[
        ['Framework do site', 'Next.js 16 (React)', 'Gratuito'],
        ['Base de dados', 'PostgreSQL (Supabase)', 'Gratuito'],
        ['Armazenamento de imagens', 'Supabase Storage', 'Gratuito'],
        ['Alojamento web', 'Vercel (plano gratuito)', 'Gratuito'],
        ['Domínio', 'cbaangola.org (Hostinger)', '~8.980 Kz/ano'],
        ['Certificado SSL (HTTPS)', "Let's Encrypt (automático)", 'Gratuito'],
    ],
    col_widths_pct=[30, 45, 25],
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_heading(doc, '4.2 Funcionalidades Entregues', level=2)

features = [
    ('Página Inicial', ' — com banner principal, avisos urgentes, matéria em destaque e publicações recentes;'),
    ('Notícias', ' — sistema completo de notícias com filtros por categoria (Notícia, Evento, Seminário, Conferência, Encontro);'),
    ('Publicações', ' — artigos, devocionais, revista EBD, com autor, resumo e conteúdo completo. 12 categorias disponíveis;'),
    ('Departamentos', ' — 9 departamentos da CBA com cards visuais e páginas de detalhe;'),
    ('Galeria de Eventos', ' — galeria fotográfica com visualizador em ecrã completo (lightbox);'),
    ('Encontre Uma Igreja', ' — 52 igrejas registadas, organizadas por 10 províncias, com pesquisa integrada;'),
    ('Quem Somos', ' — 7 secções editáveis (Visão e Valores, Historial, Declaração Doutrinária, Pacto, Baptistas, Estatuto, Plano Estratégico);'),
    ('Contactos', ' — formulário funcional que guarda mensagens na base de dados; horário de funcionamento do escritório;'),
    ('Doações', ' — sistema de doações com 3 categorias, IBAN, e envio de comprovativos;'),
    ('Página Online', ' — player de YouTube incorporado com chat ao vivo para transmissões de cultos;'),
    ('Botão Flutuante de Live', ' — botão vermelho activável pelo administrador quando há culto em directo;'),
    ('Painel Administrativo', ' — gestão completa de todo o conteúdo do site, com 13 secções (Avisos, Notícias, Publicações, Galeria, Igrejas, Departamentos, Quem Somos, Contactos, Doações, Mensagens, Comprovativos, Definições, Dashboard);'),
    ('Sistema de Upload', ' — upload de imagens e documentos PDF directamente pelo painel admin, com compressão automática;'),
    ('PWA (App Instalável)', ' — o site pode ser instalado no telemóvel como uma aplicação, com ícone no ecrã inicial;'),
    ('Navegação Mobile', ' — barra de navegação inferior fixa (estilo app) e menu hamburger com ícones SVG;'),
    ('Vídeos', ' — suporte a vídeos do YouTube e Vimeo incorporados em notícias e publicações;'),
    ('Responsivo', ' — site totalmente adaptado a telemóveis, tablets e computadores;'),
    ('SEO e Performance', ' — optimização de imagens, cache, lazy loading, e metadados para motores de busca.'),
]
for title, rest in features:
    add_bullet(doc, [
        {'text': title, 'bold': True, 'size': 10.5},
        {'text': rest, 'size': 10.5},
    ])

# ════════════════════════════════════════════
# SECTION 5: DEVOLUÇÃO DO SALDO
# ════════════════════════════════════════════
p = doc.add_paragraph()
p.add_run().add_break(WD_BREAK.PAGE)

add_heading(doc, '5. Devolução do Saldo Remanescente', level=1)

add_paragraph_with_runs(doc, [
    {'text': 'Após a aquisição do domínio '},
    {'text': 'cbaangola.org', 'bold': True},
    {'text': ' e a conclusão de todas as despesas inerentes ao projecto, verificou-se um '},
    {'text': 'saldo remanescente de 14.361,00 Kz', 'bold': True},
    {'text': ' (catorze mil, trezentos e sessenta e um Kwanzas), proveniente da diferença entre o valor depositado '
             '(100.000,00 Kz) e o custo efectivo do domínio.'}
])

add_paragraph_with_runs(doc, [
    {'text': 'O referido valor encontra-se à disposição para devolução à Convenção Baptista de Angola, através '
             'de uma das seguintes modalidades, aguardando indicação da Direcção sobre a forma preferida:'}
])

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_table_styled(
    doc,
    headers=['Modalidade', 'Detalhes'],
    rows=[
        ['Transferência Bancária', 'Transferência para conta bancária indicada pela CBA'],
        ['Entrega em Mão', 'Entrega física do valor em numerário, em local e data a acordar'],
        ['Multicaixa Express', 'Envio via Multicaixa Express para número de telefone indicado'],
    ],
    col_widths_pct=[30, 70],
)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_paragraph_with_runs(doc, [
    {'text': 'Aguarda-se feedback', 'bold': True},
    {'text': ' da Direcção sobre a modalidade preferida para que a devolução seja efectuada atempadamente.'}
])

# ════════════════════════════════════════════
# SECTION 6: CONCLUSÃO
# ════════════════════════════════════════════
add_heading(doc, '6. Conclusão', level=1)

add_paragraph_with_runs(doc, [
    {'text': 'O projecto de construção do site oficial da Convenção Baptista de Angola encontra-se '},
    {'text': 'concluído com sucesso', 'bold': True},
    {'text': ', com todas as funcionalidades planeadas implementadas, testadas e operacionais. O site está '
             'acessível ao público no endereço '},
    {'text': 'https://www.cbaangola.org', 'bold': True},
    {'text': ', com certificado de segurança SSL, painel administrativo completo, e capacidade de instalação '
             'como aplicação no telemóvel.'}
])

add_paragraph_with_runs(doc, [
    {'text': 'A gestão integral do conteúdo do site — incluindo notícias, publicações, igrejas, departamentos, '
             'avisos, galeria, definições e demais elementos — pode ser realizada autonomamente pela Direcção '
             'através do painel administrativo em '},
    {'text': 'https://www.cbaangola.org/#/admin', 'bold': True},
    {'text': ', sem necessidade de conhecimentos técnicos avançados.'}
])

add_paragraph_with_runs(doc, [
    {'text': 'Os custos de manutenção anual são minimizados, resumindo-se apenas à renovação do domínio '
             '(aproximadamente 8.980 Kz/ano), uma vez que o alojamento, a base de dados e o armazenamento de '
             'imagens são fornecidos gratuitamente pelas plataformas Vercel e Supabase.'}
])

add_paragraph_with_runs(doc, [
    {'text': 'O saldo remanescente de 14.361,00 Kz está disponível para devolução imediata, aguardando-se '
             'apenas a indicação da modalidade preferida pela Direcção.'}
])

# Signature block
doc.add_paragraph().paragraph_format.space_after = Pt(30)

sig_line = doc.add_paragraph()
sig_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
sig_line.paragraph_format.space_after = Pt(8)
add_para_border(sig_line, 'bottom', 'D4D0C8', '8')
# Force the line to render with a thin width by adding empty text
r = sig_line.add_run(' ')
set_run_font(r, size=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Hermenegildo José Pombal')
set_run_font(r, name='Cambria', size=12, bold=True, color=CBA_DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Desenvolvedor do Projecto')
set_run_font(r, name='Calibri', size=10, color=CBA_MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run('Luanda, Julho de 2026')
set_run_font(r, name='Calibri', size=10, color=CBA_MUTED)

# ─── Add footer with page numbers ───
section = doc.sections[0]
footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Relatório de Prestação de Contas — CBA  |  Julho 2026  |  Página ')
set_run_font(r, name='Calibri', size=8, color=CBA_MUTED)
# PAGE field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'PAGE \\* MERGEFORMAT'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
r2 = footer_p.add_run()
r2._r.append(fldChar1)
r2._r.append(instrText)
r2._r.append(fldChar2)
set_run_font(r2, name='Calibri', size=8, color=CBA_MUTED)

# Save
doc.save(output_path)
print(f'✅ DOCX gerado: {output_path}')
print(f'   Tamanho: {os.path.getsize(output_path) / 1024:.1f} KB')
